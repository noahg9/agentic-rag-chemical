import sys
import os

# Add the parent directory of 'work-validation' to the Python path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from meta_agent import run_meta_agent

import os
import atexit
import threading
from dotenv import load_dotenv
from langsmith import Client, traceable
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from typing import Dict, Any, List
import json

# Load environment variables
load_dotenv()

# LangSmith setup
os.environ["LANGSMITH_TRACING"] = os.getenv("LANGSMITH_TRACING", "true")
os.environ["LANGSMITH_ENDPOINT"] = os.getenv("LANGSMITH_ENDPOINT", "https://api.smith.langchain.com")
os.environ["LANGSMITH_API_KEY"] = os.getenv("LANGSMITH_API_KEY")
LANGSMITH_PROJECT = os.getenv("LANGSMITH_PROJECT", "workstations")
client = Client()

# Model selection mapping
MODEL_CHOICES = {
    "1": "azure",
    "2": "ollama",
    "3": "qwen2.5"
}

# Test queries
TEST_QUERIES = [
    "How do I handle hypophosphorous acid?",
    "What safety measures are required for unloading a tanker?",
    "What are the steps for HPLC Sulfated Analysis?",
    "How do I properly clean the reactor after use?",
    "What is the required PPE for handling sulfuric acid?"
]


def select_model() -> str:
    """Select and set the LLM model."""
    print("🔹 Choose LLM Model:")
    for key, model in MODEL_CHOICES.items():
        print(f"{key}: {model.title()}")

    choice = input("Enter 1, 2, or 3: ").strip()
    selected = MODEL_CHOICES.get(choice, "ollama")
    os.environ["ACTIVE_LLM_MODEL"] = selected
    print(f"\n✅ Using {selected} as the LLM model.")
    return selected


def print_result(query: str, result: Dict[str, Any]):
    """Print formatted results to terminal."""
    print("\n" + "=" * 80)
    print(f"Query: {query}")
    print("-" * 80)
    print(f"Combined Answer: {result['combined_answer']}")
    print(f"Confidence: {result['confidence']}%")
    print(f"Evaluation Feedback: {result['evaluation_feedback']}")
    print("\nWork Result:")
    print(json.dumps(result['work_result'], indent=2))
    print("\nSource Documents:")
    for doc in result['source_documents']:
        print(f"• {doc}")
    print("=" * 80 + "\n")


def save_results_to_docx(results: Dict[str, Any], filename: str = "meta_agent_results.docx") -> bool:
    """Save results to DOCX format."""
    try:
        doc = DocxDocument()
        doc.add_heading("Meta-Agent Routing Validation Results", 0)

        for query, details in results.items():
            doc.add_heading(f"Query: {query}", level=1)
            doc.add_paragraph(f"Combined Answer: {details['combined_answer']}")
            doc.add_paragraph(f"Confidence: {details['confidence']}%")
            doc.add_paragraph(f"Evaluation Feedback: {details['evaluation_feedback']}")
            doc.add_paragraph(f"Work Result: {str(details['work_result'])}")

            doc.add_heading("Retrieved Documents:", level=2)
            for doc_name in details['source_documents']:
                doc.add_paragraph(f"• {doc_name}", style='List Bullet')
            doc.add_page_break()

        doc.save(filename)
        print(f"✅ Results saved to DOCX: {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to save DOCX: {e}")
        return False


def save_results_to_pdf(results: Dict[str, Any], filename: str = "meta_agent_results.pdf") -> bool:
    """Save results to PDF format."""
    try:
        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        y_pos = height - 50

        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y_pos, "Meta-Agent Routing Validation Results")

        for query, details in results.items():
            y_pos = write_pdf_section(c, query, details, y_pos, width, height)

        c.save()
        print(f"✅ Results saved to PDF: {filename}")
        return True
    except Exception as e:
        print(f"❌ Failed to save PDF: {e}")
        return False


def write_pdf_section(c, query: str, details: Dict[str, Any], y: float, width: float, height: float) -> float:
    """Write a section to PDF and return new y position."""
    if y < 100:
        c.showPage()
        y = height - 50

    c.setFont("Helvetica", 12)
    c.drawString(50, y, f"Query: {query}")
    y -= 20

    for key, value in details.items():
        if key == "source_documents":
            continue
        text = f"{key.replace('_', ' ').title()}: {str(value)}"
        c.drawString(70, y, text)
        y -= 20

    if y < 100:
        c.showPage()
        y = height - 50

    return y


@traceable(project_name=LANGSMITH_PROJECT)
def process_query(query: str) -> Dict[str, Any]:
    """Process a single query with tracing."""
    result = run_meta_agent(query)
    return {
        "combined_answer": result.get("combined_answer", ""),
        "work_result": result.get("work_result", {}),
        "source_documents": [doc.get("Document Name", "Unknown")
                             for doc in result.get("source_documents", [])],
        "confidence": result.get("confidence", 0),
        "evaluation_feedback": result.get("evaluation_feedback", "No feedback")
    }


def cleanup_threads():
    """Clean up background threads."""
    for thread in threading.enumerate():
        if isinstance(thread, threading.Thread) and thread != threading.current_thread():
            try:
                if hasattr(thread, 'stop_event'):
                    thread.stop_event.set()
                thread.join(timeout=1.0)
            except Exception:
                pass


def shutdown_cleanup():
    """Clean up function to run at exit."""
    if hasattr(client, 'close'):
        client.close()

    # Clean up langsmith background threads
    import langsmith._internal._background_thread as bg_thread
    if hasattr(bg_thread, 'TRACER_THREAD'):
        try:
            bg_thread.TRACER_THREAD.stop_event.set()
            bg_thread.TRACER_THREAD.join(timeout=1.0)
        except Exception:
            pass

    cleanup_threads()


def main():
    """Main execution flow."""
    atexit.register(shutdown_cleanup)

    try:
        select_model()
        routing_results = {}

        print("\n🔄 Running work-validation tests for Meta-Agent routing...")
        for i, query in enumerate(TEST_QUERIES, 1):
            print(f"\n🔍 Processing test query {i}/{len(TEST_QUERIES)}: '{query}'")
            result = process_query(query)
            routing_results[query] = result
            print_result(query, result)
            print(f"✅ Completed query: '{query}'")

        # Save results after printing
        saved_docx = save_results_to_docx(routing_results)
        saved_pdf = save_results_to_pdf(routing_results)

        if not (saved_docx or saved_pdf):
            print("\n❌ Failed to save results to files. Check permissions and try again.")

    except Exception as e:
        print(f"❌ Error during execution: {e}")
    finally:
        shutdown_cleanup()


if __name__ == "__main__":
    main()