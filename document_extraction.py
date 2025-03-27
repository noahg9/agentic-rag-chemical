try:
    import pdfplumber
except ImportError:
    pdfplumber = None

import pandas as pd
from docx import Document as DocxDocument

from llm import get_llm

# Initialize the LLM model
llm_model = get_llm("qwen2.5")  # Change the model choice as needed


def extract_nested_tables(table):
    extracted_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            if cell.tables:
                nested_data = extract_nested_tables(cell.tables[0])
                extracted_data.extend(nested_data)
            else:
                row_data.append(cell.text.strip())
        if row_data:
            extracted_data.append(row_data)
    return extracted_data


def query_table(tables, keyword):
    """
    Searches for a specific keyword in the tables and returns relevant data.
    """
    results = []
    for table in tables:
        if keyword in table.values:
            row_idx, col_idx = (table == keyword).to_numpy().nonzero()
            for r, c in zip(row_idx, col_idx):
                if c < len(table.columns) - 1:  # Ensure there's data next to the keyword
                    results.append((table.iloc[r, c], table.iloc[r, c + 1]))  # (Parameter, Value)
    return results


def extract_table_data_from_docx(docx_path):
    """
    Extracts structured table data, including nested tables, from a DOCX file.
    """
    doc = DocxDocument(docx_path)
    tables_data = []

    for table in doc.tables:
        table_rows = extract_nested_tables(table)  # Extract nested table data
        if table_rows:
            tables_data.append(pd.DataFrame(table_rows))  # Convert to DataFrame

    return tables_data


def extract_full_text_from_docx(docx_path: str) -> str:
    """
    Extracts full text from a DOCX file, including paragraphs and tables (both as plain and structured text).

    Returns:
        str: The complete text extracted from the DOCX file.
    """
    doc = DocxDocument(docx_path)
    text_elements = []

    # Extract paragraphs.
    for para in doc.paragraphs:
        if para.text.strip():
            text_elements.append(para.text.strip())

    # Extract tables as plain text.
    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                table_lines.append(" | ".join(row_data))
        if table_lines:
            text_elements.append("TABLE DATA:\n" + "\n".join(table_lines))

    # Append structured table data.
    structured_tables = extract_table_data_from_docx(docx_path)
    if structured_tables:
        table_text = "\n\n".join(
            [table.to_string(index=False) for table in structured_tables])  # Convert tables to readable string
        text_elements.append("STRUCTURED TABLE DATA:\n" + table_text)

    return "\n".join(text_elements)


def extract_table_data_from_pdf(pdf_path: str):
    """
    Extracts structured table data from a PDF file using pdfplumber.

    Returns:
        List: A list of tables extracted from the PDF, each represented as structured data if headers exist.
    """
    if pdfplumber is None:
        print("pdfplumber not installed. Skipping PDF table extraction.")
        return []
    tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if not table:
                    continue
                if len(table) > 1:
                    headers = table[0]
                    table_data = []
                    for row in table[1:]:
                        row_dict = {}
                        for i, cell in enumerate(row):
                            header = headers[i] if i < len(headers) and headers[i] else f"Column{i + 1}"
                            row_dict[header] = cell
                        table_data.append(row_dict)
                    tables.append(table_data)
                else:
                    tables.append(table)
    print("Extracted PDF table data:", tables)
    return tables


def extract_full_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts full text from a PDF file using pdfplumber.

    Returns:
        str: The complete text extracted from the PDF.
    """
    if pdfplumber is None:
        print("pdfplumber not installed. Skipping PDF text extraction.")
        return ""
    with pdfplumber.open(pdf_path) as pdf:
        text_elements = []
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_elements.append(page_text)
        return "\n".join(text_elements)


def llm_query_sop(documents, user_query, llm_model):
    """
    Uses an LLM to answer user queries based on both document text and structured tables.
    """
    context = ""

    for doc in documents:
        context += f"\nDOCUMENT: {doc.metadata.get('source', 'Unknown')}\n"
        context += doc.page_content  # Add text content

        if "tables" in doc.metadata:
            for table in doc.metadata["tables"]:
                context += "\nTABLE DATA:\n"
                context += table.to_string(index=False)  # Convert DataFrame to readable format

    # Construct an LLM prompt
    prompt = (
        "You are an expert assistant analyzing standard operating procedures.\n"
        "Given the document content and structured table data below, answer the following question accurately:\n"
        f"QUESTION: {user_query}\n\n"
        "DOCUMENT DATA:\n"
        f"{context}\n"
        "Please provide a concise and accurate response."
    )

    return llm_model.run(prompt)
