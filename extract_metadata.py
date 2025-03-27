
from docx import Document as DocxDocument

import fitz  # PyMuPDF
import os
import re

from pyparsing import matchPreviousLiteral


def extract_metadata_from_pdf(pdf_path: str) -> dict:
    """
    Extracts metadata from PDFs, including author, keywords, structural details,
    as well as the document's version and date.
    """
    doc = fitz.open(pdf_path)
    standard_metadata = doc.metadata
    first_page = doc.load_page(0)
    first_page_text = first_page.get_text("text")
    doc.close()

    custom_metadata = {
        "Document Name": os.path.basename(pdf_path).replace(".pdf", ""),
        "Step Number": None,
        "Author": standard_metadata.get("author", "Unknown"),
        "Created Date": standard_metadata.get("creationDate", "Unknown"),
        "Keywords": standard_metadata.get("keywords", "").split(", "),
        "Process Type": None,
        "Safety Measures": [],
        "Version": None,
        "Date": None,
    }

    patterns = {
        "Step Number": re.compile(r"Step\s+(\d+)", re.IGNORECASE),
        "Process Type": re.compile(r"Process Type[:\s]*(.*)", re.IGNORECASE),
        "Safety Measures": re.compile(r"Safety Measures[:\s]*(.*)", re.IGNORECASE),
        "Version": re.compile(r"Version[:\s]*(\S+)", re.IGNORECASE),
        "Date": re.compile(r"Date\s*[:]\s*([\d/]+)", re.IGNORECASE),
    }

    for key, pattern in patterns.items():
        match = pattern.search(first_page_text)
        if match:
            custom_metadata[key] = match.group(1).strip()

    metadata = {**standard_metadata, **custom_metadata}
    return metadata

def filter_metadata(metadata: dict) -> dict:
    filtered = {}
    for k, v in metadata.items():
        if isinstance(v, (str, int, float, bool)):
            filtered[k] = v
        elif v is None:
            filtered[k] = ""
        elif isinstance(v, list):
            filtered[k] = ", ".join(map(str, v))  # Convert lists to strings
        else:
            filtered[k] = str(v)
    return filtered



def extract_metadata_from_docx(docx_path: str) -> dict:
    print("[DEBUG] extract_metadata_from_docx() function started.")
    doc = DocxDocument(docx_path)

    print(f"[DEBUG] Found {len(doc.tables)} tables in the document.")

    metadata = {
        "Document Name": os.path.basename(docx_path).replace(".docx", ""),
        "Review": "",
        "Date": "",
    }

    # Flexible regex to match multi-line cases
    patterns = {
        "Review": re.compile(r"Review[:\s\n]*([\d]+)", re.IGNORECASE),
        "Date": re.compile(r"Date[:\s\n]*([\d]{2}/[\d]{2}/[\d]{4})", re.IGNORECASE),
    }

    def clean_text(text):
        """Normalize text by removing extra spaces and special characters."""
        return " ".join(text.replace("\xa0", " ").split()).strip()

    # Extract text from tables
    for table_idx, table in enumerate(doc.tables):
        print(f"[DEBUG] Processing Table {table_idx + 1}")
        table_text = []

        for row in table.rows:
            row_text = []  # Store text from each row
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                row_text.append(cell_text)

            # Join row text into a single line (handles multi-line cell cases)
            table_text.append(" | ".join(row_text))

        # Join the entire table text to allow regex to work across multi-line data
        table_text = "\n".join(table_text)

        print(f"\n[DEBUG] Table {table_idx + 1} Extracted Text:\n{table_text}\n")

        # Apply regex search on the full extracted table text
        for key, pattern in patterns.items():
            if not metadata[key]:  # Only extract if not found yet
                match = pattern.search(table_text)
                if match:
                    metadata[key] = match.group(1).strip()
                    print(f"[MATCH] Found {key}: {metadata[key]}")

    return metadata
