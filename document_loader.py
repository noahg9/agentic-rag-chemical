from typing import List
from langchain_community.document_loaders import PyPDFLoader
from extract_metadata import extract_metadata_from_pdf, extract_metadata_from_docx, filter_metadata
from document_extraction import extract_full_text_from_pdf, extract_full_text_from_docx, extract_table_data_from_pdf, \
    extract_table_data_from_docx
from langchain.schema import Document
from docx import Document as DocxDocument
import os
import pdfplumber
from PIL import Image, UnidentifiedImageError
from io import BytesIO
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_community.llms import Ollama
from document_extraction import extract_full_text_from_pdf, extract_full_text_from_docx, extract_table_data_from_docx, \
    query_table
from extract_metadata import extract_metadata_from_pdf, extract_metadata_from_docx
from retriever import filter_complex_metadata


#
# from pdf2image import convert_from_path
# from PIL import Image
# import pytesseract
#
# def extract_images_from_pdf(pdf_path):
#     """Extract images from a PDF file using pdf2image."""
#     images = convert_from_path(pdf_path)
#     return images
#
# def preprocess_image(image):
#     """Enhance image for better OCR."""
#     image = image.convert('L')  # Convert to grayscale
#     image = image.resize((image.width * 2, image.height * 2))  # Scale up
#     return image
#
#
#
# def ocr_with_llava(image):
#     """Runs OCR on an image using LLaVA 7B with a structured prompt."""
#     try:
#         # Convert PIL image to base64
#         import base64
#         import io
#
#         buffered = io.BytesIO()
#         image.save(buffered, format="JPEG")
#         img_str = base64.b64encode(buffered.getvalue()).decode()
#
#         llm = Ollama(model="llava:latest", temperature=0.2)
#         prompt = "Extract and structure any readable text from this image. Ensure clarity for technical work instructions."
#
#         # Pass the base64 encoded image
#         response = llm.invoke(input=prompt, images=[img_str])
#
#         print("OCR Response:", response)
#         return response
#     except Exception as e:
#         print(f"OCR failed for image: {str(e)}")
#         return "OCR FAILED"
#
#
# # document_loader.py
# def extract_text_from_images(images):
#     extracted_texts = []
#     for i, image in enumerate(images):
#         image = preprocess_image(image)
#         text = pytesseract.image_to_string(image)
#         if text.strip():
#             extracted_texts.append(f"IMAGE {i}: {text}")
#     return "\n".join(extracted_texts)



def load_work_instructions(directory_path):
    """Load PDF work instructions with metadata, tables, and image OCR."""
    work_instruction_docs = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".pdf"):
            metadata = extract_metadata_from_pdf(file_path)
            filtered_metadata = filter_metadata(metadata)
            text = extract_full_text_from_pdf(file_path)
            table_data = extract_table_data_from_pdf(file_path)
            # images = extract_images_from_pdf(file_path)
            # image_text = extract_text_from_images(images) if images else ""

            doc_content = text
            # if image_text:
            #     doc_content += "\n\nIMAGE OCR TEXT: " + image_text
            if table_data:
                doc_content += "\n\nTABLE DATA: " + str(table_data)

            doc = Document(page_content=doc_content, metadata=filtered_metadata)
            doc.metadata.update(metadata)
            work_instruction_docs.append(Document(page_content=text , metadata=metadata))
            return work_instruction_docs


def load_sop_documents(directory_path):
    """
    Load DOCX SOPs with metadata and structured table data.
    """
    sop_docs = []

    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith(".docx"):
            metadata = extract_metadata_from_docx(file_path)
            text = extract_full_text_from_docx(file_path)
            tables = extract_table_data_from_docx(file_path)  # Get structured tables

            doc = Document(page_content=text, metadata=metadata)
            doc.metadata.update(metadata)

            # Ensure structured tables are stored inside metadata
            if tables:
                doc.metadata["tables"] = [table.to_string(index=False) for table in tables] # Convert to string

            cleaned_metadata = filter_complex_metadata(metadata)
            doc = Document(page_content=text, metadata=cleaned_metadata)
            sop_docs.append(doc)

    return sop_docs


def query_sop_table(documents, keyword):
    """
    Queries all SOP documents for a specific parameter in the extracted tables.

    Returns:
        A dictionary of document names and matching values.
    """
    results = {}

    for doc in documents:
        if "tables" in doc.metadata:
            for table in doc.metadata["tables"]:
                result = query_table([table], keyword)
                if result:
                    results[doc.metadata.get("source", "Unknown Document")] = result

    return results
